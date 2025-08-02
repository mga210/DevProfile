#!/usr/bin/env python3
"""
Script to create a proper DOCX resume file using python-docx
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new Run object and add the text
    new_run = OxmlElement('w:r')

    # Set the run's style to hyperlink
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_job_entry(doc, title, company, dates, bullets):
    """Add a job entry with consistent formatting"""
    # Job title and company
    p = doc.add_paragraph()
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)
    
    p.add_run(f" | {company}")
    
    # Dates (right-aligned)
    date_p = doc.add_paragraph(dates)
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.space_after = Pt(6)
    
    # Bullet points
    for bullet in bullets:
        bullet_p = doc.add_paragraph(bullet, style='List Bullet')
        bullet_p.space_after = Pt(3)

def create_docx_resume():
    """Create a professional DOCX resume"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name and contact
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_run = header.add_run("Miguel A. Gonzalez Almonte")
    name_run.font.size = Pt(18)
    name_run.bold = True
    
    contact = doc.add_paragraph("Plano, TX • 787-367-9843 • mgonzalez869@gmail.com")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(6)
    
    # Create paragraph with clickable links
    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add hyperlinks
    add_hyperlink(links, "https://www.linkedin.com/in/miguel-gonzalez-8a389791", "LinkedIn")
    links.add_run(" • ")
    add_hyperlink(links, "https://mga210.github.io/DevProfile/", "Portfolio")
    links.add_run(" • ")
    add_hyperlink(links, "https://github.com/mga210", "GitHub")
    
    links.space_after = Pt(12)
    
    # Professional tagline
    tagline = doc.add_paragraph("AI Systems Builder | GPT-Powered Workflow Architect | Operational Intelligence Technologist")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.runs[0]
    tagline_run.font.size = Pt(12)
    tagline_run.bold = True
    tagline.space_after = Pt(12)
    
    # Professional Summary
    add_heading_with_style(doc, "PROFESSIONAL SUMMARY", 1)
    
    summary_text = """AI systems builder with deep roots in operations leadership and a self-driven path into software design, automation, and GPT agent architecture. I've built and deployed intelligent tools that replaced spreadsheets, reduced operational delays, and brought logic-based coordination into live operational environments. My systems are not experiments — they are active solutions used in production to enforce lifecycle integrity, audit state transitions, and support non-technical users.

Though I've never held a formal software title, I've architected AI assistants, Python dashboards, and decision frameworks that reflect enterprise-grade thinking and field-tested pragmatism. My work lives at the intersection of operational intelligence and systems clarity — where good logic can save time, prevent errors, and build trust.

Now seeking roles where I can continue designing agent-powered workflows, internal tools, and smart coordination systems — especially in environments that value practical intelligence, not just pedigree."""
    
    doc.add_paragraph(summary_text).space_after = Pt(12)
    
    # Skills Summary
    add_heading_with_style(doc, "SKILLS SUMMARY", 1)
    
    skills_data = [
        ("AI & Workflow Intelligence", [
            "GPT Agent Design, Prompt Architecture, Logic Scaffolding",
            "Recursive Flows, Task Routing Systems, Lifecycle Modeling",
            "Role-Based Interaction Flows, Instructional UX for Non-Technical Users"
        ]),
        ("Python Systems Development", [
            "Python 3 (Modular Scripting, DTO Structures, API Basics)",
            "Pandas (Data Structuring, Filtering, Export Pipelines)",
            "PySide6 and Tkinter (GUI Architecture, Interaction Flows)",
            "FastAPI (Lightweight API Services)",
            "SQLite and Supabase (Layered DB Use, State Tracing, Portability)"
        ]),
        ("Automation & Data Tools", [
            "Excel (Advanced Formulas, VBA, Macros, Conditional Logic)",
            "Power BI (Custom Dashboards, Workflow Reporting)",
            "CMMS Data Structuring, Delay Tracking Logic, Heatmaps and Export Systems"
        ])
    ]
    
    for category, skills in skills_data:
        cat_p = doc.add_paragraph()
        cat_run = cat_p.add_run(category)
        cat_run.bold = True
        cat_run.font.size = Pt(11)
        
        for skill in skills:
            skill_p = doc.add_paragraph(f"• {skill}")
            skill_p.space_after = Pt(3)
    
    # Professional Experience
    add_heading_with_style(doc, "PROFESSIONAL EXPERIENCE", 1)
    
    # Job 1
    add_job_entry(doc, "Service Maintenance Manager", "MAA – Dallas, TX", "Jun 2023 – Present", [
        "Led service operations across 3 multifamily properties while designing and deploying the Make Ready Digital Board (DMRB) — a logic-based AI tool used live to coordinate unit readiness",
        "Replaced manual spreadsheets with a state-resolved Python system using DTOs, task templates, and lifecycle enforcement",
        "Built audit-safe, role-scoped task flows with offline queueing, conflict detection, and automatic readiness locking",
        "Reduced unit turnover time from 13–20 days to 7 through system-led coordination and real-time visibility",
        "Integrated Python (Pandas, PySide6, FastAPI), SQLite, and Supabase to enable full-stack functionality"
    ])
    
    # Job 2
    add_job_entry(doc, "Service Manager", "RPM Living – Dallas, TX", "May 2022 – Jun 2023", [
        "Directed daily service operations at a high-volume multifamily property, overseeing technician workflows, vendor schedules, and turnover timelines",
        "Built and deployed custom Excel dashboards to reduce admin friction, improve task tracking, and align team focus",
        "Applied Agile-style planning methods to improve technician coverage and reduce backlog",
        "Created SOPs and structured vendor workflows to streamline unit turnover"
    ])
    
    # Job 3
    add_job_entry(doc, "Operations Assistant → Kitchen Manager", "Universal Studios – Orlando, FL", "2009 – 2018", [
        "Progressed from operations assistant to leading two full-service kitchens during high-volume park operations",
        "Participated in engineering launch teams for new venues and attractions",
        "Led cross-kitchen menu rollout projects, coordinating timing, staff training, and guest flow readiness",
        "Developed early awareness of system bottlenecks, team handoffs, and operations logic under pressure"
    ])
    
    # Education & Certifications
    add_heading_with_style(doc, "EDUCATION & CERTIFICATIONS", 1)
    
    edu_p = doc.add_paragraph()
    edu_run = edu_p.add_run("Bachelor of Business Administration in Computer Information Systems")
    edu_run.bold = True
    edu_p.add_run("\nAna G. Méndez University – Carolina, PR (In Progress)")
    edu_p.space_after = Pt(6)
    
    cert_p = doc.add_paragraph()
    cert_run = cert_p.add_run("Completed Certifications:")
    cert_run.bold = True
    
    certifications = [
        "Python for Everybody – University of Michigan / Coursera (2025)",
        "Python 3 – Intermediate Track (2025)",
        "Google Project Management Certificate – Coursera (2025)",
        "EPA Section 608 Certification – HVAC Systems (2018)"
    ]
    
    for cert in certifications:
        cert_bullet = doc.add_paragraph(f"• {cert}")
        cert_bullet.space_after = Pt(3)
    
    # Key Projects
    add_heading_with_style(doc, "KEY PROJECTS", 1)
    
    projects = [
        ("Make Ready Digital Board (DMRB)", "AI-Powered Task Lifecycle System", [
            "Replaced spreadsheets with a logic-resolved unit coordination engine deployed across 3 properties",
            "Reduced turnover from 13–20 days to 7 with role-based access, task gating, and offline queueing"
        ]),
        ("System Pilot", "GPT-Powered Architecture Strategist", [
            "Designed a GPT assistant that guides users from raw product ideas into full system blueprints",
            "Enforced architectural planning through modular dialogue and logic scaffolding"
        ]),
        ("Blueprint Buddy", "Modular GPT Instruction Builder", [
            "Built a scalable system to translate user input into structured, validated GPT instructions",
            "Used by product leads and toolmakers to deploy prompt systems with embedded logic flows"
        ])
    ]
    
    for project_name, project_type, bullets in projects:
        proj_p = doc.add_paragraph()
        proj_name_run = proj_p.add_run(project_name)
        proj_name_run.bold = True
        proj_name_run.font.size = Pt(11)
        
        proj_type_p = doc.add_paragraph(project_type)
        proj_type_p.runs[0].italic = True
        proj_type_p.space_after = Pt(3)
        
        for bullet in bullets:
            bullet_p = doc.add_paragraph(f"• {bullet}")
            bullet_p.space_after = Pt(3)
    
    # Save the document
    doc.save('resume.docx')
    print("Professional DOCX resume created successfully: resume.docx")

if __name__ == "__main__":
    create_docx_resume()